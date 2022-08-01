import docx,os
str1 = r"""已知数列$\{a_n\}$满足$a_1=1$, $S_n=\dfrac{(n+1)a_n}2$($n\in \mathbf{N}^*$), 求通项$a_n$的表达式."""
str2 = r"""$2005$年$1$月$6$日, 我国人口总数为$13$亿, 称该天为``中国人口$13$亿日'', 如果$2005$年$1$月$6$日后我国人口的年自然增长率保持在$0.6\%$, 问到哪一年我国人口总数将超过$14$亿?
"""
str3 = r"""(1) 若约定$\rho>0$, $0\le \theta <2\pi$, 试写出图中$A$、$B$、$C$、$D$、$E$各点的极坐标$(\rho,\theta)$;\\
(2) 若约定$\rho<0$, $0\le\theta <2\pi$, 试写出图中$A$、$B$、$C$、$D$、$E$各点的极坐标$(\rho,\theta)$.
\begin{center}
\begin{tikzpicture}[>=latex,scale = 0.5]
\foreach \i in {0,15,...,345} {\draw (0,0) -- (\i:5.1);};
\foreach \i in {1,2,...,5} {\draw (0,0) circle (\i); \draw (\i,0) node [below right] {$\i$};};
\draw (30:6) node {$\dfrac \pi 6$};
\draw (60:6) node {$\dfrac \pi 3$};
\draw (90:6) node {$\dfrac \pi 2$};
\draw (120:6) node {$\dfrac {2\pi} 3$};
\draw (150:6) node {$\dfrac {5\pi} 6$};
\draw (180:6) node {$\pi$};
\draw (210:6) node {$\dfrac {7\pi} 6$};
\draw (240:6) node {$\dfrac {4\pi} 3$};
\draw (270:6) node {$\dfrac {3\pi} 2$};
\draw (300:6) node {$\dfrac {5\pi} 3$};
\draw (330:6) node {$\dfrac {11\pi} 6$};
\filldraw (90:3) circle (0.1) node [above right] {$A$};
\filldraw (150:4) circle (0.1) node [above] {$B$};
\filldraw (15:5) circle (0.1) node [right] {$C$};
\filldraw (210:3.5) circle (0.1) node [above] {$D$};
\filldraw (315:1) circle (0.1) node [below] {$E$};
\end{tikzpicture}
\end{center}"""
problems = [str1,str2,str3]
doc = docx.Document()
doc.add_paragraph("一、解答题")
count = 0
for p in problems:
    with open("eq.tex","w",encoding="utf8") as f:
        f.write(p)
    count += 1
    os.system("xelatex equations.tex")
    os.system("pdftocairo equations.pdf -png -singlefile -r 300 pics/eq")
    doc.add_paragraph(str(count)+". ")
    doc.add_picture("pics/eq.png")
doc.save("test02.docx")
    